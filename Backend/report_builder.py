"""
report_builder.py
─────────────────
Builds a professional Word (.docx) DDR report from the AI-generated JSON.

Responsibilities:
  - Create branded UrbanRoof DDR document
  - Place all 7 required sections
  - Embed relevant images under correct area sections
  - Handle missing data gracefully (write "Not Available")
  - Output clean, client-friendly Word document
"""

import io
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from pdf_extractor import PDFExtractResult, ExtractedImage


# ─────────────────────────────────────────────────────────────
#  BRAND COLORS
# ─────────────────────────────────────────────────────────────

COLOR_ORANGE  = RGBColor(0xF6, 0xAD, 0x55)   # UrbanRoof orange
COLOR_DARK    = RGBColor(0x1A, 0x1D, 0x2E)   # Dark navy
COLOR_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_GREY    = RGBColor(0x71, 0x80, 0x96)
COLOR_RED     = RGBColor(0xE5, 0x3E, 0x3E)   # High severity
COLOR_YELLOW  = RGBColor(0xD6, 0x9E, 0x28)   # Moderate severity
COLOR_GREEN   = RGBColor(0x38, 0xA1, 0x69)   # Low severity
COLOR_BLACK   = RGBColor(0x1A, 0x20, 0x2C)


# ─────────────────────────────────────────────────────────────
#  HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_horizontal_line(doc, color_hex="F6AD55"):
    """Add a colored horizontal rule paragraph."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def _severity_color(level: str) -> RGBColor:
    """Return color based on severity level."""
    level = level.lower() if level else ""
    if "critical" in level or "high" in level:
        return COLOR_RED
    elif "moderate" in level:
        return COLOR_YELLOW
    elif "low" in level:
        return COLOR_GREEN
    return COLOR_GREY


def _get_value(data: dict, key: str, fallback: str = "Not Available") -> str:
    """Safely get a value from dict, return fallback if missing/empty."""
    val = data.get(key, fallback)
    if not val or str(val).strip() == "":
        return fallback
    return str(val).strip()


# ─────────────────────────────────────────────────────────────
#  REPORT BUILDER CLASS
# ─────────────────────────────────────────────────────────────

class ReportBuilder:
    """
    Builds the DDR Word document.

    Usage:
        builder  = ReportBuilder()
        docx_buf = builder.build(
                       ddr_data,
                       inspection_result,
                       thermal_result,
                       property_info
                   )
    """

    def build(
        self,
        ddr_data: dict,
        inspection_result: PDFExtractResult,
        thermal_result: PDFExtractResult,
        property_info: dict,
    ) -> bytes:
        """
        Main entry point.
        Returns Word document as bytes.
        """
        doc = Document()

        # Set page margins
        for section in doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(2.5)
            section.right_margin  = Cm(2.5)

        # Build image lookup: label -> ExtractedImage
        image_map = {}
        for img in inspection_result.images + thermal_result.images:
            image_map[img.label] = img

        # ── Build all sections ───────────────────────────────
        self._add_cover_page(doc, property_info)
        self._add_disclaimer(doc)
        self._add_property_summary(doc, property_info)
        self._add_issue_summary(doc, ddr_data)
        self._add_area_observations(doc, ddr_data, image_map)
        self._add_root_cause(doc, ddr_data)
        self._add_severity_assessment(doc, ddr_data)
        self._add_recommended_actions(doc, ddr_data)
        self._add_additional_notes(doc, ddr_data)
        self._add_missing_info(doc, ddr_data)
        self._add_footer_note(doc)

        # Save to bytes
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    # ─────────────────────────────────────────────────────────
    #  COVER PAGE
    # ─────────────────────────────────────────────────────────

    def _add_cover_page(self, doc: Document, property_info: dict):
        # Company name
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("UrbanRoof Private Limited")
        run.bold      = True
        run.font.size = Pt(22)
        run.font.color.rgb = COLOR_ORANGE

        # Tagline
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run("Building Health Diagnosis Specialists")
        r2.font.size = Pt(11)
        r2.font.color.rgb = COLOR_GREY
        r2.italic = True

        doc.add_paragraph()
        _add_horizontal_line(doc)
        doc.add_paragraph()

        # Report title
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run("DETAILED DIAGNOSIS REPORT")
        r3.bold      = True
        r3.font.size = Pt(26)
        r3.font.color.rgb = COLOR_DARK

        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.add_run("(DDR)")
        r4.bold = True
        r4.font.size = Pt(16)
        r4.font.color.rgb = COLOR_ORANGE

        doc.add_paragraph()
        doc.add_paragraph()

        # Property details table
        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        details = [
            ("Prepared For",   property_info.get("customer_name", "Not Available")),
            ("Property",       property_info.get("property_address", "Not Available")),
            ("Property Type",  property_info.get("property_type", "Not Available")),
            ("Inspection Date",str(property_info.get("inspection_date", "Not Available"))),
            ("Inspected By",   property_info.get("inspected_by", "Not Available")),
            ("Report Date",    str(date.today())),
        ]

        for label, value in details:
            row = table.add_row()
            # Label cell
            label_cell = row.cells[0]
            _set_cell_bg(label_cell, "F6AD55")
            lp = label_cell.paragraphs[0]
            lr = lp.add_run(label)
            lr.bold = True
            lr.font.color.rgb = COLOR_DARK
            lr.font.size = Pt(11)

            # Value cell
            value_cell = row.cells[1]
            vp = value_cell.paragraphs[0]
            vr = vp.add_run(value)
            vr.font.size = Pt(11)
            vr.font.color.rgb = COLOR_BLACK

        doc.add_page_break()

    # ─────────────────────────────────────────────────────────
    #  DISCLAIMER
    # ─────────────────────────────────────────────────────────

    def _add_disclaimer(self, doc: Document):
        self._section_heading(doc, "Data and Information Disclaimer")

        disclaimer_text = (
            "This property inspection report is not an exhaustive inspection of the structure, "
            "systems, or components. The inspection may not reveal all deficiencies. "
            "A health check-up helps to reduce some of the risk involved in the property, "
            "but it cannot eliminate these risks, nor can the inspection anticipate future events "
            "or changes in performance due to changes in use or occupancy.\n\n"
            "An inspection addresses only those components and conditions that are present, "
            "visible, and accessible at the time of the inspection. This report is based on "
            "visual and non-destructive examination only."
        )

        p = doc.add_paragraph(disclaimer_text)
        p.runs[0].font.size = Pt(10)
        p.runs[0].font.color.rgb = COLOR_GREY
        p.runs[0].italic = True
        doc.add_paragraph()

    # ─────────────────────────────────────────────────────────
    #  PROPERTY SUMMARY TABLE
    # ─────────────────────────────────────────────────────────

    def _add_property_summary(self, doc: Document, property_info: dict):
        self._section_heading(doc, "Section 1 · Property & Inspection Details")

        table = doc.add_table(rows=0, cols=2)
        table.style = "Table Grid"

        rows_data = [
            ("Customer Name",       property_info.get("customer_name",    "Not Available")),
            ("Property Address",    property_info.get("property_address", "Not Available")),
            ("Property Type",       property_info.get("property_type",    "Not Available")),
            ("Number of Floors",    str(property_info.get("floors",       "Not Available"))),
            ("Property Age",        f"{property_info.get('property_age', 'Not Available')} years"),
            ("Inspection Date",     str(property_info.get("inspection_date", "Not Available"))),
            ("Inspected By",        property_info.get("inspected_by",     "Not Available")),
            ("Previous Audit Done", property_info.get("prev_audit",       "Not Available")),
            ("Previous Repairs",    property_info.get("prev_repair",      "Not Available")),
        ]

        for label, value in rows_data:
            row   = table.add_row()
            lc    = row.cells[0]
            vc    = row.cells[1]
            _set_cell_bg(lc, "2D3748")
            lp    = lc.paragraphs[0]
            lr    = lp.add_run(label)
            lr.bold = True
            lr.font.color.rgb = COLOR_ORANGE
            lr.font.size = Pt(10)
            vp    = vc.paragraphs[0]
            vr    = vp.add_run(value)
            vr.font.size = Pt(10)

        doc.add_paragraph()

    # ─────────────────────────────────────────────────────────
    #  SECTION 1: PROPERTY ISSUE SUMMARY
    # ─────────────────────────────────────────────────────────

    def _add_issue_summary(self, doc: Document, ddr_data: dict):
        self._section_heading(doc, "Section 2 · Property Issue Summary")

        summary = _get_value(ddr_data, "property_issue_summary")

        # Highlighted summary box
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]
        _set_cell_bg(cell, "FFF3CD")
        p = cell.paragraphs[0]
        r = p.add_run(summary)
        r.font.size = Pt(11)
        r.font.color.rgb = COLOR_BLACK

        doc.add_paragraph()

    # ─────────────────────────────────────────────────────────
    #  SECTION 2: AREA-WISE OBSERVATIONS
    # ─────────────────────────────────────────────────────────

    def _add_area_observations(
        self,
        doc: Document,
        ddr_data: dict,
        image_map: dict,
    ):
        self._section_heading(doc, "Section 3 · Area-wise Observations")

        observations = ddr_data.get("area_wise_observations", [])

        if not observations:
            p = doc.add_paragraph("Not Available — No area observations found in the documents.")
            p.runs[0].font.color.rgb = COLOR_GREY
            return

        for obs in observations:
            area_name = _get_value(obs, "area_name", "Unknown Area")

            # Area sub-heading
            ph = doc.add_paragraph()
            rh = ph.add_run(f"  {area_name}  ")
            rh.bold = True
            rh.font.size = Pt(13)
            rh.font.color.rgb = COLOR_WHITE
            rh.font.highlight_color = None
            # Use shaded paragraph via table trick
            tbl = doc.add_table(rows=1, cols=1)
            tbl.style = "Table Grid"
            cell = tbl.rows[0].cells[0]
            _set_cell_bg(cell, "2D3748")
            cp = cell.paragraphs[0]
            cr = cp.add_run(f"▶  {area_name}")
            cr.bold = True
            cr.font.size = Pt(12)
            cr.font.color.rgb = COLOR_ORANGE

            # Observation table
            obs_table = doc.add_table(rows=0, cols=2)
            obs_table.style = "Table Grid"

            obs_rows = [
                ("Negative Side\n(Impacted Area)",  _get_value(obs, "negative_side")),
                ("Positive Side\n(Source Area)",    _get_value(obs, "positive_side")),
                ("Thermal Reading",                 _get_value(obs, "thermal_reading")),
            ]

            for lbl, val in obs_rows:
                row = obs_table.add_row()
                lc  = row.cells[0]
                vc  = row.cells[1]
                _set_cell_bg(lc, "F7FAFC")
                lp  = lc.paragraphs[0]
                lr  = lp.add_run(lbl)
                lr.bold = True
                lr.font.size = Pt(10)
                lr.font.color.rgb = COLOR_DARK
                vp  = vc.paragraphs[0]
                vr  = vp.add_run(val)
                vr.font.size = Pt(10)

            doc.add_paragraph()

            # ── Images for this area ─────────────────────────
            image_labels = obs.get("images", [])
            placed_count = 0

            for label in image_labels:
                img_obj = image_map.get(label)
                if img_obj is None:
                    # Image label given but not found
                    p_miss = doc.add_paragraph()
                    r_miss = p_miss.add_run(f"[Image Not Available: {label}]")
                    r_miss.italic = True
                    r_miss.font.color.rgb = COLOR_GREY
                    r_miss.font.size = Pt(9)
                    continue

                try:
                    img_stream = io.BytesIO(img_obj.image_bytes)
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    run_img = p_img.add_run()
                    run_img.add_picture(img_stream, width=Inches(2.8))

                    # Image caption
                    p_cap = doc.add_paragraph()
                    r_cap = p_cap.add_run(
                        f"Figure: {label} | "
                        f"Source: {img_obj.source.capitalize()} Report | "
                        f"Page {img_obj.page_number}"
                    )
                    r_cap.italic = True
                    r_cap.font.size = Pt(8)
                    r_cap.font.color.rgb = COLOR_GREY
                    placed_count += 1

                except Exception:
                    p_err = doc.add_paragraph()
                    r_err = p_err.add_run(f"[Could not embed image: {label}]")
                    r_err.italic = True
                    r_err.font.size = Pt(9)
                    r_err.font.color.rgb = COLOR_GREY

            if image_labels and placed_count == 0:
                p_na = doc.add_paragraph()
                r_na = p_na.add_run("Image Not Available for this area.")
                r_na.italic = True
                r_na.font.size = Pt(9)
                r_na.font.color.rgb = COLOR_GREY

            _add_horizontal_line(doc, "E2E8F0")
            doc.add_paragraph()

    # ─────────────────────────────────────────────────────────
    #  SECTION 3: PROBABLE ROOT CAUSE
    # ─────────────────────────────────────────────────────────

    def _add_root_cause(self, doc: Document, ddr_data: dict):
        self._section_heading(doc, "Section 4 · Probable Root Cause")

        root_cause = _get_value(ddr_data, "probable_root_cause")

        p = doc.add_paragraph()
        r = p.add_run(root_cause)
        r.font.size = Pt(11)

        # Conflicts
        conflicts = ddr_data.get("conflicts_found", [])
        if conflicts:
            doc.add_paragraph()
            ph = doc.add_paragraph()
            rh = ph.add_run("⚠️  Conflicts Found Between Reports:")
            rh.bold = True
            rh.font.color.rgb = COLOR_RED
            rh.font.size = Pt(11)

            for conflict in conflicts:
                pc = doc.add_paragraph(style="List Bullet")
                rc = pc.add_run(str(conflict))
                rc.font.size = Pt(10)
                rc.font.color.rgb = COLOR_RED

        doc.add_paragraph()

    # ─────────────────────────────────────────────────────────
    #  SECTION 4: SEVERITY ASSESSMENT
    # ─────────────────────────────────────────────────────────

    def _add_severity_assessment(self, doc: Document, ddr_data: dict):
        self._section_heading(doc, "Section 5 · Severity Assessment")

        severity = ddr_data.get("severity_assessment", {})
        overall  = _get_value(severity, "overall_level")
        reasoning= _get_value(severity, "reasoning")

        # Overall severity badge
        sev_table = doc.add_table(rows=1, cols=2)
        sev_table.style = "Table Grid"
        lc = sev_table.rows[0].cells[0]
        vc = sev_table.rows[0].cells[1]
        _set_cell_bg(lc, "2D3748")
        lp = lc.paragraphs[0]
        lr = lp.add_run("Overall Severity Level")
        lr.bold = True
        lr.font.color.rgb = COLOR_ORANGE
        lr.font.size = Pt(12)

        sev_color = _severity_color(overall)
        vp = vc.paragraphs[0]
        vr = vp.add_run(overall.upper())
        vr.bold = True
        vr.font.size = Pt(14)
        vr.font.color.rgb = sev_color

        doc.add_paragraph()

        # Reasoning
        ph = doc.add_paragraph()
        rh = ph.add_run("Reasoning:")
        rh.bold = True
        rh.font.size = Pt(11)

        pr = doc.add_paragraph(reasoning)
        pr.runs[0].font.size = Pt(11)
        doc.add_paragraph()

        # Area-wise severity table
        area_severities = severity.get("area_wise_severity", [])
        if area_severities:
            ph2 = doc.add_paragraph()
            rh2 = ph2.add_run("Area-wise Severity:")
            rh2.bold = True
            rh2.font.size = Pt(11)

            tbl = doc.add_table(rows=1, cols=3)
            tbl.style = "Table Grid"

            # Header row
            hdr = tbl.rows[0]
            headers = ["Area", "Severity", "Reason"]
            for i, h in enumerate(headers):
                _set_cell_bg(hdr.cells[i], "F6AD55")
                p  = hdr.cells[i].paragraphs[0]
                r  = p.add_run(h)
                r.bold = True
                r.font.color.rgb = COLOR_DARK
                r.font.size = Pt(10)

            for item in area_severities:
                row  = tbl.add_row()
                area = _get_value(item, "area")
                sev  = _get_value(item, "severity")
                rsn  = _get_value(item, "reason")

                row.cells[0].paragraphs[0].add_run(area).font.size = Pt(10)
                sev_run = row.cells[1].paragraphs[0].add_run(sev)
                sev_run.font.color.rgb = _severity_color(sev)
                sev_run.bold = True
                sev_run.font.size = Pt(10)
                row.cells[2].paragraphs[0].add_run(rsn).font.size = Pt(10)

        doc.add_paragraph()

    # ─────────────────────────────────────────────────────────
    #  SECTION 5: RECOMMENDED ACTIONS
    # ─────────────────────────────────────────────────────────

    def _add_recommended_actions(self, doc: Document, ddr_data: dict):
        self._section_heading(doc, "Section 6 · Recommended Actions")

        actions = ddr_data.get("recommended_actions", [])

        if not actions:
            p = doc.add_paragraph("Not Available — No recommended actions found.")
            p.runs[0].font.color.rgb = COLOR_GREY
            return

        # Group by priority
        priority_order = ["Immediate", "Short Term", "Long Term"]
        grouped = {"Immediate": [], "Short Term": [], "Long Term": [], "Other": []}

        for action in actions:
            priority = action.get("priority", "Other")
            if priority in grouped:
                grouped[priority].append(action)
            else:
                grouped["Other"].append(action)

        priority_colors = {
            "Immediate":  "E53E3E",
            "Short Term": "D69E28",
            "Long Term":  "38A169",
            "Other":      "718096",
        }

        for priority in priority_order:
            items = grouped.get(priority, [])
            if not items:
                continue

            # Priority heading
            tbl = doc.add_table(rows=1, cols=1)
            tbl.style = "Table Grid"
            _set_cell_bg(tbl.rows[0].cells[0], priority_colors[priority])
            ph = tbl.rows[0].cells[0].paragraphs[0]
            rh = ph.add_run(f"🔴  {priority.upper()} ACTION REQUIRED" if priority == "Immediate"
                           else f"🟡  {priority.upper()}" if priority == "Short Term"
                           else f"🟢  {priority.upper()}")
            rh.bold = True
            rh.font.color.rgb = COLOR_WHITE
            rh.font.size = Pt(11)

            for item in items:
                p  = doc.add_paragraph(style="List Bullet")
                r  = p.add_run(
                    f"{_get_value(item, 'action')}  "
                    f"[Area: {_get_value(item, 'area')}]"
                )
                r.font.size = Pt(10)

            doc.add_paragraph()

    # ─────────────────────────────────────────────────────────
    #  SECTION 6: ADDITIONAL NOTES
    # ─────────────────────────────────────────────────────────

    def _add_additional_notes(self, doc: Document, ddr_data: dict):
        self._section_heading(doc, "Section 7 · Additional Notes")

        notes = _get_value(ddr_data, "additional_notes")
        p = doc.add_paragraph(notes)
        p.runs[0].font.size = Pt(11)
        doc.add_paragraph()

    # ─────────────────────────────────────────────────────────
    #  SECTION 7: MISSING / UNCLEAR INFO
    # ─────────────────────────────────────────────────────────

    def _add_missing_info(self, doc: Document, ddr_data: dict):
        self._section_heading(doc, "Section 8 · Missing or Unclear Information")

        missing = ddr_data.get("missing_or_unclear_information", [])

        if not missing:
            p = doc.add_paragraph("All required information was available in the provided documents.")
            p.runs[0].font.color.rgb = COLOR_GREEN
            p.runs[0].font.size = Pt(11)
            return

        for item in missing:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(str(item) if item else "Not Available")
            r.font.size = Pt(10)
            r.font.color.rgb = COLOR_GREY

        doc.add_paragraph()

    # ─────────────────────────────────────────────────────────
    #  FOOTER NOTE
    # ─────────────────────────────────────────────────────────

    def _add_footer_note(self, doc: Document):
        _add_horizontal_line(doc)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(
            f"This report was generated by DiagnoAI · UrbanRoof Private Limited · "
            f"www.urbanroof.in · +91-8925-805-805 · "
            f"Report Date: {date.today()}"
        )
        r.font.size = Pt(9)
        r.font.color.rgb = COLOR_GREY
        r.italic = True

    # ─────────────────────────────────────────────────────────
    #  REUSABLE SECTION HEADING
    # ─────────────────────────────────────────────────────────

    def _section_heading(self, doc: Document, title: str):
        """Add a styled section heading."""
        doc.add_paragraph()
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(14)
        r.font.color.rgb = COLOR_ORANGE
        _add_horizontal_line(doc)
        doc.add_paragraph()